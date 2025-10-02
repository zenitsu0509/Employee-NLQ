from __future__ import annotations

import io
import json
import mimetypes
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Sequence

import numpy as np
import pandas as pd
from backend.api.services.job_tracker import JobProgress, JobStatus
from backend.api.services.vector_store import DocumentChunk, VectorStore


SUPPORTED_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/csv": "csv"
}

mimetypes.add_type("application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx")
mimetypes.add_type("text/csv", ".csv")


@dataclass
class ProcessedDocument:
    document_id: str
    chunks: List[DocumentChunk]


class DocumentProcessor:
    """Service that ingests and indexes unstructured documents."""

    def __init__(self, vector_store: VectorStore, model: "SentenceTransformer", batch_size: int) -> None:
        self._vector_store = vector_store
        self._model = model
        self._batch_size = batch_size

    def process_documents(self, file_paths: Iterable[Path], job: JobProgress | None = None) -> List[ProcessedDocument]:
        processed_docs: List[ProcessedDocument] = []
        paths = list(file_paths)
        total_files = len(paths)

        if job:
            job.total = total_files
            job.status = JobStatus.RUNNING

        for index, path in enumerate(paths, start=1):
            text_chunks = self._extract_chunks(path)
            embeddings = self._embed_chunks(text_chunks)
            doc_chunks = [
                DocumentChunk(
                    chunk_id=f"{path.stem}-{i}",
                    document_id=path.name,
                    content=chunk,
                    metadata={"path": str(path), "order": i}
                )
                for i, chunk in enumerate(text_chunks)
            ]
            if embeddings.size:
                self._vector_store.add(embeddings, doc_chunks)
            processed_docs.append(ProcessedDocument(document_id=path.name, chunks=doc_chunks))

            if job:
                job.processed = index
                job.message = f"Processed {path.name}"

        if job:
            job.status = JobStatus.COMPLETED
            job.message = f"Indexed {total_files} document(s)"

        return processed_docs

    def dynamic_chunking(self, content: str, doc_type: str) -> List[str]:
        """Intelligent chunking heuristics based on document type."""
        if doc_type == "resume":
            return self._chunk_by_sections(content, ["skills", "experience", "projects", "education"], max_chars=800)
        if doc_type == "contract":
            return self._chunk_by_sections(content, ["section", "clause", "article"], max_chars=1200)
        if doc_type == "review":
            return self._chunk_paragraphs(content, max_chars=600)
        if doc_type == "table":
            return self._chunk_csv_rows(content, max_rows_per_chunk=10)
        return self._chunk_sentences(content, max_chars=800)

    # Extraction helpers

    def _extract_chunks(self, path: Path) -> List[str]:
        mime_type, _ = mimetypes.guess_type(path)
        doc_type = SUPPORTED_TYPES.get(mime_type, "unknown")
        text = self._read_file(path, doc_type)
        inferred_type = self._infer_document_type(text, path)
        return self.dynamic_chunking(text, inferred_type)

    def _read_file(self, path: Path, doc_type: str) -> str:
        if doc_type == "pdf":
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        if doc_type == "docx":
            from docx import Document  # type: ignore

            doc = Document(str(path))
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        if doc_type == "csv":
            df = pd.read_csv(path)
            return df.to_csv(index=False)
        if path.suffix.lower() in {".json", ".jsonl"}:
            return self._read_json(path)
        return path.read_text(encoding="utf-8", errors="ignore")

    def _read_json(self, path: Path) -> str:
        text = path.read_text(encoding="utf-8", errors="ignore")
        try:
            if path.suffix.lower() == ".jsonl":
                return "\n".join(json.dumps(json.loads(line)) for line in text.splitlines() if line)
            return json.dumps(json.loads(text), indent=2)
        except json.JSONDecodeError:
            return text

    def _infer_document_type(self, text: str, path: Path) -> str:
        lower = text.lower()
        if "resume" in lower or "objective" in lower or "skills" in lower:
            return "resume"
        if "agreement" in lower or "clause" in lower or "party" in lower:
            return "contract"
        if "performance" in lower and "review" in lower:
            return "review"
        if path.suffix.lower() in {".csv", ".tsv"}:
            return "table"
        return "generic"

    # Chunking helpers

    def _chunk_by_sections(self, content: str, keywords: Sequence[str], max_chars: int) -> List[str]:
        sections: Dict[str, List[str]] = {keyword: [] for keyword in keywords}
        current_section = "other"
        buffer: List[str] = []

        for line in content.splitlines():
            lower_line = line.lower()
            for keyword in keywords:
                if keyword in lower_line:
                    if buffer:
                        sections.setdefault(current_section, []).append("\n".join(buffer).strip())
                        buffer = []
                    current_section = keyword
                    break
            buffer.append(line)

        if buffer:
            sections.setdefault(current_section, []).append("\n".join(buffer).strip())

        return self._merge_chunks([chunk for sec_chunks in sections.values() for chunk in sec_chunks], max_chars)

    def _chunk_paragraphs(self, content: str, max_chars: int) -> List[str]:
        paragraphs = [paragraph.strip() for paragraph in content.split("\n\n") if paragraph.strip()]
        return self._merge_chunks(paragraphs, max_chars)

    def _chunk_sentences(self, content: str, max_chars: int) -> List[str]:
        reader = io.StringIO(content)
        sentences: List[str] = []
        current = []
        current_len = 0
        for line in reader:
            for sentence in self._sentence_tokenize(line.strip()):
                length = len(sentence)
                if current_len + length > max_chars and current:
                    sentences.append(" ".join(current).strip())
                    current = [sentence]
                    current_len = length
                else:
                    current.append(sentence)
                    current_len += length
        if current:
            sentences.append(" ".join(current).strip())
        return sentences

    def _merge_chunks(self, chunks: List[str], max_chars: int) -> List[str]:
        merged: List[str] = []
        buffer = ""
        for chunk in chunks:
            if len(buffer) + len(chunk) <= max_chars:
                buffer = f"{buffer}\n{chunk}" if buffer else chunk
            else:
                if buffer:
                    merged.append(buffer.strip())
                buffer = chunk
        if buffer:
            merged.append(buffer.strip())
        return merged

    def _sentence_tokenize(self, text: str) -> List[str]:
        delimiters = {".", "?", "!"}
        sentence = []
        sentences: List[str] = []
        for char in text:
            sentence.append(char)
            if char in delimiters:
                sentences.append("".join(sentence).strip())
                sentence = []
        if sentence:
            sentences.append("".join(sentence).strip())
        return sentences

    def _chunk_csv_rows(self, content: str, max_rows_per_chunk: int = 10) -> List[str]:
        """Chunk CSV content by preserving header + batches of rows."""
        lines = content.strip().splitlines()
        if not lines:
            return []
        
        header = lines[0] if lines else ""
        data_rows = lines[1:]
        
        chunks: List[str] = []
        for start in range(0, len(data_rows), max_rows_per_chunk):
            batch_rows = data_rows[start:start + max_rows_per_chunk]
            chunk = "\n".join([header] + batch_rows)
            chunks.append(chunk)
        
        return chunks if chunks else [content]

    def _embed_chunks(self, chunks: List[str]) -> np.ndarray:
        embeddings: List[np.ndarray] = []
        for start in range(0, len(chunks), self._batch_size):
            batch = chunks[start:start + self._batch_size]
            if not batch:
                continue
            batch_embeddings = self._model.encode(batch, show_progress_bar=False)
            embeddings.append(np.array(batch_embeddings))
        return np.vstack(embeddings) if embeddings else np.empty((0, self._vector_store.dimension))
